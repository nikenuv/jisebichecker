from docx import Document
import docx
import re
from docx.shared import Pt  # Import Pt for font size comparison
import matplotlib.pyplot as plt
from collections import OrderedDict
from docx2pdf import convert
from pdf2docx import parse
import random
import pythoncom

def get_title(directory):
    doc = Document(directory)
    
    # Assume the title is in the first paragraph of the document
    title_paragraph = doc.paragraphs[0]
    
    # Extract the text from the paragraph
    title = title_paragraph.text
    
    return title

def get_author(directory):
    doc = Document(directory)
      
    # Assume the authors are in the second paragraph of the document
    authors_paragraph = doc.paragraphs[1]
    authors = authors_paragraph.text
    
    return authors

def cek_abstrak(directory):
    return []

def cek_abstrak_sequential(directory):
    
    # Load an existing document.
    doc = Document(directory)
    regex = [
    r"^Abstract",
    r"^Background:",
    r"^Objective[s]?:",
    r"^Method[s]?:",
    r"^Result[s]?:",
    r"^Conclusion[s]?:",
    r"^Keywords:",
    r"^Article history:",
    ]



    # Check if the document contains a table.
    if doc.tables:
        # Get the first table in the document.
        first_table = doc.tables[0]

        # Print the first table.
        # print("First Table:")
        for row in first_table.rows:
            # print(row.text)
            for cell in row.cells:
                teks = cell.text
                print(cell.text)
                break
        teks = first_table.rows[0].cells[0].text
        print("HERERE", teks)
    else:
        print("The document does not contain a table.")


    # print(teks)
    hasil=[0] * (8)
    total=[0] * (100)
    tokat = 0

    if 'teks' not in locals():
        return False, False, False, False, False, False, False, False, False

    
    for line_number, line in enumerate(teks.split('\n'), 1):
        for i in range (8):
            if re.match(regex[i], line.strip()):
                # print(f"Line {line_number}: {line.strip()} matches pattern {regex[i]}")
                if i == 6:
                    keyword = line.strip()
                hasil[i]= line_number
        word_count = len(line.strip().split())  # Menghitung jumlah kata dalam baris
        # print(f"kata: {word_count}")
        total[line_number]= word_count

    # for i in range (8):
    #     print(hasil[i])

    for i in range (hasil[6]-1):
        tokat = tokat + total[i]
        
    print(f"Jumlah kata: {tokat}")

    if tokat >= 150 and tokat <=300:
        tokats = True
    else:
        tokats = False

    if 'keyword' in locals():
        if keyword.count(',')+1 >=3 and keyword.count(',')+1 <= 6:
            jumkey = True
        else:
            jumkey = False
    else:
        jumkey = False

    is_sorted = all(hasil[i] <= hasil[i + 1] for i in range(len(hasil) - 1))



    if hasil[1] != 0:
        background = True
    else:
        background = False
    if hasil[2] != 0:
        objectives = True
    else:
        objectives = False
    if hasil[3] != 0:
        methods = True
    else:
        methods = False
    if hasil[4] != 0:
        result = True
    else:
        result = False
    if hasil[5] != 0:
        conclusion = True
    else:
        conclusion = False
    if hasil[6] != 0:
        keyword = True
    else:
        keyword = False

    return background, objectives, methods, result, conclusion, keyword, is_sorted, jumkey, tokats

    

def cek_footer(directory):
    different_first_page_header_footer = True
    doc = Document(directory)
    regex = [
        r"ISSN 2443-2555 \(online\) 2598-6333 \(print\) © 2023 The Authors\. Published by Universitas Airlangga\.\s+This is an open access article under the CC BY license \([^)]*\)",
        r'doi: http://dx\.doi\.org/10\.20473/jisebi\.[0-9]+\.[0-9]+\.[0-9]+-[0-9]+',
    ]
    regek = [
        r"^ISSN 2443-2555 \(online\) 2598-6333 \(print\) © 20XX The Authors\. Published by Universitas Airlangga\.\s+This is an open access article under the CC BY license \(http://creativecommons\.org/licenses/by/4\.0/\)$",
        r'^doi: http://dx\.doi\.org/10\.20473/jisebi\.X\.X\.pp-pp$',
    ]
    line = [None] * 2
    formatfooter = True
    font_size = True
    i = 0
    for section in doc.sections:
        header = section.first_page_footer
        if header is not None:
            # print("Footer ditemukan di bagian ini:")
            for paragraph in header.paragraphs:
                # print(paragraph.text)
                if re.match(regek[i], paragraph.text):
                    formatfooter = formatfooter
                    # print("Format sesuai")

                else:
                    formatfooter = False
                    # print('Format tidak sesuai')
                for run in paragraph.runs:
                    line[i] = run.font
                    break
                # line[i]=paragraph.run.font
                if i == 1:
                    break
                i=i+1
                if i == 5:
                    break
    
    if all(obj is not None for obj in line):
        i = 0
        for x in range(1):
            if line[x].size != Pt(8):
                font_size = False
    else:
        font_size = False

    result = {
        "format": formatfooter,
        "font_size": font_size
    }
    return result

def cek_header_1(directory):
    different_first_page_header_footer = True
    doc = Document(directory)
    regex = [
        r'^Journal of$',
        r'^Information Systems Engineering$',
        r'^and Business Intelligence$',
        r'^Vol\.[0-9]+, No\.[0-9]+, [A-Za-z]+ [0-9]+$',
        r'^Available online at: http://e-journal\.unair\.ac\.id/index\.php/JISEBI$'
    ]
    regek = [
        r'^Journal of$',
        r'^Information Systems Engineering$',
        r'^and Business Intelligence$',
        r"^Vol\.X, No\.X, February/June/October XXXX$",
        r'Available online at: http://e-journal\.unair\.ac\.id/index\.php/JISEBI$'
    ]
    line = [None] * 5
    hasil = 0
    i = 0

    formatheader = True
    font_size_title = True
    font_size_details = True
    font_family = True
    


    for section in doc.sections:
        header = section.first_page_header

        if header is not None:
            # print("Header ditemukan di bagian ini:")
            for paragraph in header.paragraphs:
                # print(paragraph.text)
                if re.match(regex[i], paragraph.text) == True :
                    hasil = hasil + 1
                    formatheader = False
                #     print("Header Oke")
                # else:
                #     hasil = hasil + 1
                #     print('Format tidak sesuai')
                for run in paragraph.runs:
                    line[i] = run.font
                    # print(run.font.size)
                    break
                # line[i]=paragraph.run.font
                if i == 4:
                    break
                i=i+1
                if i == 5:
                    break
        else:
            result = {
                'format': False,
                'font_size_title': False,
                'font_size_details': False,
                'font_family': False
            }
            return result

    if all(obj is not None for obj in line):
        i = 0
        for x in range(4):
            if x <= 2:
                if line[x].size != Pt(14):
                    hasil = hasil + 1
                    font_size_title = False
            else:
                if line[x].size != Pt(8):
                    hasil = hasil + 1
                    font_size_details = False
            if line[x].name != 'Californian FB':
                hasil = hasil + 1
                font_family = False
            if x == 5:
                break
    else:
        print("NONEE")
        result = {
            'format': False,
            'font_size_title': False,
            'font_size_details': False,
            'font_family': False
        }
        return result

    result = {
        'format': formatheader,
        'font_size_title': font_size_title,
        'font_size_details': font_size_details,
        'font_family': font_family
    }
    return result

def cek_header_normal(directory):
    different_first_page_header_footer = True

    doc = Document(directory)
    regex = [
        r'^[A-Za-z]+, [A-Za-z]+, [A-Za-z]+, [A-Za-z]+, [A-Za-z]+, & [A-Za-z]+\s+Journal Information Systems Engineering and Business Intelligence, [0-9]+, [0-9]+ \([0-9]+\), [0-9]+-[0-9]+',
        r'^[A-Za-z]+, [A-Za-z]+, [A-Za-z]+, [A-Za-z]+, & [A-Za-z]+\s+Journal Information Systems Engineering and Business Intelligence, [0-9]+, [0-9]+ \([0-9]+\), [0-9]+-[0-9]+',
        r'^[A-Za-z]+, [A-Za-z]+, [A-Za-z]+, & [A-Za-z]+\s+Journal Information Systems Engineering and Business Intelligence, [0-9]+, [0-9]+ \([0-9]+\), [0-9]+-[0-9]+',
        r'^[A-Za-z]+, [A-Za-z]+, & [A-Za-z]+\s+Journal Information Systems Engineering and Business Intelligence, [0-9]+, [0-9]+ \([0-9]+\), [0-9]+-[0-9]+',
        r'^[A-Za-z]+ & [A-Za-z]+\s+Journal Information Systems Engineering and Business Intelligence, [0-9]+, [0-9]+ \([0-9]+\), [0-9]+-[0-9]+',
        r'^[A-Za-z]+\s+Journal Information Systems Engineering and Business Intelligence, [0-9]+, [0-9]+ \([0-9]+\), [0-9]+-[0-9]+',
    ]
    regek = r"^Last Name Author1, Last Name Author2, & Last Name Author3\s+Journal of Information Systems Engineering and Business Intelligence, XXXX, X \(X\), pp-pp$"
    line = [None] * 5
    hasil = 0
    i = 0
    formatHeader = True
    font_size = True
    for section in doc.sections:
        header = section.header
        if header is not None:
            # print("Header ditemukan di bagian ini:")
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    if run.font.size != Pt(8):
                        font_size = False
                # print(paragraph.text)
                for i in range(6):
                    if re.match(regek, paragraph.text):
                        break
                    if i == 5:
                        formatHeader = False

    result = {
        "format": formatHeader,
        "font_size": font_size,
    }

    return result

def cek_title(docx_file_path):

    size = False
    bold = False
    doc = docx.Document(docx_file_path)
    # Periksa apakah ada paragraf pertama dalam dokumen
    if len(doc.paragraphs) > 0:
        first_paragraph = doc.paragraphs[0]
        # print(first_paragraph.text.strip())
        for run in first_paragraph.runs:
            if run.font.size == Pt(18):
                size = True
            else:
                if  run.font.size == None:
                    size = True
                else:
                    size = False
            if run.bold == "False":
                bold = False
            else:
                bold = True
        return size, bold       
        # return bool(first_paragraph)  # Mengembalikan True jika paragraf pertama memiliki teks
    
    return False  # Mengembalikan False jika dokumen kosong atau tidak memiliki paragraf
   
def cek_author(docx_file_path):
    size = None
    bold = None
    doc = docx.Document(docx_file_path)
    # Periksa apakah ada paragraf pertama dalam dokumen
    if len(doc.paragraphs) > 0:
        first_paragraph = doc.paragraphs[1]
        # print(first_paragraph.text.strip())
        for run in first_paragraph.runs:
            if  run.font.size == Pt(11):
                size = True
            else:
                if  run.font.size == None:
                    size = True
                else:
                    size = False
            if run.bold == "False":
                bold = False
            else:
                bold = True
        return size, bold          
        
    
    return False

def cek_litrev(directory):
    doc = Document(directory)

    regex = [
    r"^Literature  Review$",
    ]
    regek = [
    r"^[A-Za-z0-9]+\. Literature  Review$"
    ]
    hasil=[0] * (1)

    litrev = False

    # Masukan Ke String Teks
    line_number = 1
    for paragraph in doc.paragraphs:
        lines = paragraph.text.split('\n')
        for line in lines:
            teks = line
            # print(f"Line {line_number}: {line}")
            for i in range(1):
                if re.match(regex[i], teks.strip()):
                    hasil[i]=line_number
            line_number += 1

    if hasil[0] == 0:
        line_number = 1
        for paragraph in doc.paragraphs:
            lines = paragraph.text.split('\n')
            for line in lines:
                teks = line
                # print(f"Line {line_number}: {line}")
                for i in range(1):
                    if re.match(regek[i], teks.strip()):
                        hasil[i]=line_number
                line_number += 1

    if hasil[0] != 0:
        litrev = True

    return litrev

def cek_struktur(directory):
    doc = Document(directory)

    regex = [
    r"^Introduction$",
    #    r"^Literature  Review$",
    r"^Methods?$",
    r"^Results?$",
    r"^Discussions?$",
    r"^Conclusions?$",
    #    r"^References$",
    ]
    regek = [
    r"^[A-Za-z0-9]+\. Introduction?$",
    #    r"^Literature  Review$",
    r"^[A-Za-z0-9]+\. Methods?$",
    r"^[A-Za-z0-9]+\. Results?$",
    r"^[A-Za-z0-9]+\. Discussions?$",
    r"^[A-Za-z0-9]+\. Conclusions?$",
    #    r"^References$",
    ]
    hasil=[0] * (5)

    # Masukan Ke String Teks
    line_number = 1
    for paragraph in doc.paragraphs:
        lines = paragraph.text.split('\n')
        for line in lines:
            teks = line
            # print(f"Line {line_number}: {line}")
            for i in range(5):
                if re.match(regex[i], teks.strip(), re.IGNORECASE):
                    hasil[i]=line_number
            line_number += 1

    if hasil[0] == 0 and hasil[4] == 0:
        line_number = 1
        for paragraph in doc.paragraphs:
            lines = paragraph.text.split('\n')
            for line in lines:
                teks = line
                # print(f"Line {line_number}: {line}")
                for i in range(5):
                    if re.match(regek[i], teks.strip(), re.IGNORECASE):
                        hasil[i]=line_number
                line_number += 1

    introduction = method = result = discussion = conclusion = sort = False

    for i in range (5):
        if i == 0:
            if hasil[i] != 0:
                introduction = True
        elif i == 1:
            if hasil[i] != 0:
                method = True
        elif i == 2:
            if hasil[i] != 0:
                result = True
        elif i == 3:
            if hasil[i] != 0:
                discussion = True
        elif i == 4:
            if hasil[i] != 0:
                conclusion = True

    is_sorted = all(hasil[i] <= hasil[i + 1] for i in range(len(hasil) - 1))

    if hasil[0] != 0 and hasil[4] != 0:
        if is_sorted:
            sort = True
        elif hasil[1] == 0:
            hasilx = [0] * (4)
            hasilx[0] = hasil[0]
            hasilx[1] = hasil[2]
            hasilx[2] = hasil[3]
            hasilx[3] = hasil[4]
            is_sorted = all(hasilx[i] <= hasilx[i + 1] for i in range(len(hasilx) - 1))
            if is_sorted:
                sort = True
    result = {
        "introduction": introduction,
        "method": method,
        "result": result,
        "discussion": discussion,
        "conclusion": conclusion,
        "sort": sort
    }

    return result

def find_references(doc):
    references = set()  # Menggunakan set untuk menghindari duplikasi

    # Membaca teks dari dokumen DOCX
    doc_text = []
    for paragraph in doc.paragraphs:
        doc_text.append(paragraph.text)
        if "REFERENCES" in paragraph.text:
            break
        elif "References" in paragraph.text:
            break
        elif "ReferenceS" in paragraph.text:
            break
        elif "Reference" in paragraph.text:
            break
        elif "REFERENCE" in paragraph.text:
            break
        else:
            continue

    # Menggabungkan teks menjadi satu string
    doc_text = " ".join(doc_text)
    # print(doc)

    # Mencari semua referensi yang cocok dengan pola IEEE
    # Pola ini mencocokkan nomor referensi dalam kurung siku [x]
    references = list(re.findall(r'\[[0-9]+\]', doc_text))
    references = references+ list(re.findall(r'\[[0-9]+,[0-9]+\]', doc_text))
    result = []
    for item in references:
        if ',' in item:
            result.extend(item.strip('[]').split(','))
        else:
            result.append(item.strip('[]'))
    while '0' in result:
        result.remove('0')
    result = list(OrderedDict.fromkeys(result))
    
    return sorted(result)

def rep(doc):
    references = set()  # Menggunakan set untuk menghindari duplikasi

    # Membaca teks dari dokumen DOCX
    doc_text = []
    for paragraph in doc.paragraphs:
        doc_text.append(paragraph.text)

    # Menggabungkan teks menjadi satu string
    doc_text = " ".join(doc_text)

    # Mencari semua referensi yang cocok dengan pola "Ref. [x]" atau "reference [x]" (di mana x adalah angka)
    ab = set(re.findall(r'(?:Ref\.|reference)\s*\[(\d+)\]', doc_text, re.IGNORECASE))

    return ab

def total(doc):
    references_section = False
    references_text = ""
    num_lines = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # print(text)
            
        # Cari awal bagian daftar pustaka
        if "REFERENCES" in text:
            references_section = True
            continue
        elif "References" in text:
            references_section = True
            continue
        elif "ReferenceS" in text:
            references_section = True
            continue
        elif "Reference" in text:
            references_section = True
            continue
        elif "REFERENCE" in text:
            references_section = True
            continue

       # Keluar dari bagian daftar pustaka jika menemukan teks lain setelahnya
        if references_section and not text:
            if num_lines > 0:
                break
            # elif num_lines == 0:
            #     break
            continue

        # Jika baris memiliki teks, tambahkan ke teks referensi dan jumlah baris
        if references_section and text:
            if num_lines > 0:
                references_text += "\n"  # Tambahkan baris baru sebelum menambahkan teks
            references_text += text
            num_lines += 1

        # num_lines = len(references_text.splitlines())
            
        # num_lines = num_lines -1
    return num_lines

def totalnew(doc):
    totalx = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Mencari pola dalam teks sel tabel
                text = cell.text
                matches = re.findall(r'\[\d+\]', text)

                if matches:
                    totalx.extend(matches)

    return totalx

def topdf(doc):
    pythoncom.CoInitialize()
    pdf = doc.replace(".docx", ".pdf")
    convert(doc, pdf)
    return pdf

def todocx(pdf):
    doc = pdf.replace(".pdf", "-new.docx")
    parse(pdf, doc)
    return doc




# Nama file dokumen DOCX yang ingin Anda periksa


def cek_reference(directory):
    # Membuka dokumen DOCX
    doc = docx.Document(directory)

    # Menemukan referensi dalam dokumen
    references= find_references(doc)

    # format = rep(doc)
    totals = total(doc)


    #print(references)
    num_references = len(references)
    #print(f"Total {num_references} referensi ditemukan.")
    #print(f"Total {totals} referensi di bagian references ditemukan.")

    jumref = formatref = kepakesemua = False

    if totals <= 1:
        doc = docx.Document(todocx(topdf(directory)))
        
        references= find_references(doc)
        num_references = len(references)

        totals = len(totalnew(doc))

        print(references)
        if num_references <= 1:
            num_references = len(references)

    if totals >= 20:
        jumref = True
    if num_references != 0:
        formatref = True
    if num_references == totals:
        kepakesemua = True
    
    result = {
        'format': formatref,
        'count': jumref,
        'used': kepakesemua
    }

    return result

def is_valid_table(table, doc, table_index):

    # Get the two paragraphs before the table
    index_before_table = table._element.getparent().index(table._element) - 1
    if index_before_table < 0:
        return False  # Not enough paragraphs before the table

    paragraphs_before_table = doc.paragraphs[index_before_table-5:index_before_table]
    text_before_table = " ".join([paragraph.text for paragraph in paragraphs_before_table])

    # Check if the text before the table contains "TABLE [NUMBER]"
    expected_text = f"TABLE {table_index}"  # Assuming tables are 0-indexed
    if expected_text not in text_before_table:
        return False

    '''
    # Check if only horizontal lines are used within the table
    for row in table.rows:
        for cell in row.cells:
            # Check if the cell has only horizontal borders
            borders = cell._element.xpath('.//w:tcBorders')
            if (
                borders and
                borders[0].xpath('@w:val', namespaces=nsmap)[0] != 'single' or
                borders[0].xpath('@w:vertical', namespaces=nsmap)[0] != 'nil' or
                borders[0].xpath('@w:horizontal', namespaces=nsmap)[0] != 'single'
            ):
                print('rong')
                return False
'''
    return True
    

def cek_tabel(directory):
    doc = Document(directory)

    for i, table in enumerate(doc.tables[1:]):  # Start from the second table
        print(i)
        print(table)
        if not is_valid_table(table, doc, i + 1):  # Adjust the index to account for skipping the first table
            return False

    return True

def cek_figure(file_path):

    # Load the DOCX document
    doc = Document(file_path)

    # Variable to track whether the document follows the guidelines
    follows_guidelines = True

    # Iterate through each paragraph in the document
    fig_number = 0
    for i, paragraph in enumerate(doc.paragraphs):
        # Check if the paragraph contains an image
        if 'Graphic' in paragraph._p.xml:
            fig_number += 1
            print(i)
            print("Figure found")
            print(doc.paragraphs[i+1].text)
            expected_text = f"Fig "
            if expected_text not in doc.paragraphs[i+1].text and fig_number != 1:
                print("False")
                follows_guidelines = False            # Check if the paragraph has a caption with an Arabic numeral (Fig. 1)
            '''
            if "Fig." in doc.paragraphs[i+1].text:
                caption_parts = paragraph.text.split(" ")
                if len(caption_parts) >= 3 and caption_parts[2].isdigit() and caption_parts[2] != "0":
                    print(f"Image caption found: {paragraph.text}")
                else:
                    print(f"Invalid image caption: {doc.paragraphs[i+1].text}")
                    follows_guidelines = False
            else:
                print(f"Image without a caption: {paragraph.text}")
                follows_guidelines = False '''

    return follows_guidelines
